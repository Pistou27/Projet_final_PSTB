"""
Ingestion des documents et création de l'index vectoriel FAISS
"""
import os
import sys
from typing import List, Dict
from pathlib import Path
import pandas as pd
from docx import Document

# LangChain imports
from langchain_community.document_loaders import (
    PyPDFLoader, 
    TextLoader,
    CSVLoader
)
from langchain.text_splitter import RecursiveCharacterTextSplitter
# Try the new HuggingFace embeddings first, fallback to community version
try:
    from langchain_huggingface import HuggingFaceEmbeddings
except ImportError:
    from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.schema import Document as LangChainDocument

# Local imports
from config import Config
from utils_pdf import create_page_metadata, PDFPageExtractor
from utils import Logger

class DocumentIngester:
    """Gestionnaire d'ingestion des documents"""
    
    def __init__(self):
        self.config = Config
        self.embeddings = None
        self.vectorstore = None
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.config.CHUNK_SIZE,
            chunk_overlap=self.config.CHUNK_OVERLAP,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
    
    def _init_embeddings(self):
        """Initialiser le modèle d'embeddings"""
        Logger.loading(f"Chargement du modèle d'embeddings: {self.config.EMBEDDING_MODEL}")
        self.embeddings = HuggingFaceEmbeddings(
            model_name=self.config.EMBEDDING_MODEL,
            model_kwargs={'device': 'cpu'},
            encode_kwargs={'normalize_embeddings': True}
        )
    
    def load_pdf_documents(self, file_path: str) -> List[LangChainDocument]:
        """Charger un document PDF avec métadonnées de page"""
        documents = []
        
        try:
            # Utiliser notre extracteur personnalisé pour les métadonnées
            extractor = PDFPageExtractor(file_path)
            pages_content = extractor.get_all_pages_content()
            
            for page_num, content in pages_content:
                if content.strip():
                    # Créer un document avec métadonnées de page
                    metadata = {
                        "source": os.path.basename(file_path),
                        "file_path": file_path,
                        "file_type": "pdf",
                        "page": page_num,
                        "pages": [page_num]
                    }
                    
                    doc = LangChainDocument(
                        page_content=content,
                        metadata=metadata
                    )
                    documents.append(doc)
            
            Logger.success(f"PDF chargé: {os.path.basename(file_path)} ({len(documents)} pages)")
            
        except Exception as e:
            Logger.error(f"Erreur lors du chargement PDF {file_path}: {e}")
        
        return documents
    
    def load_text_documents(self, file_path: str) -> List[LangChainDocument]:
        """Charger un document texte"""
        documents = []
        
        try:
            loader = TextLoader(file_path, encoding='utf-8')
            docs = loader.load()
            
            for doc in docs:
                doc.metadata.update({
                    "source": os.path.basename(file_path),
                    "file_path": file_path,
                    "file_type": "txt"
                })
                documents.append(doc)
            
            Logger.success(f"TXT chargé: {os.path.basename(file_path)}")
            
        except Exception as e:
            Logger.error(f"Erreur lors du chargement TXT {file_path}: {e}")
        
        return documents
    
    def load_docx_documents(self, file_path: str) -> List[LangChainDocument]:
        """Charger un document Word"""
        documents = []
        
        try:
            doc = Document(file_path)
            content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            if content.strip():
                metadata = {
                    "source": os.path.basename(file_path),
                    "file_path": file_path,
                    "file_type": "docx"
                }
                
                lang_doc = LangChainDocument(
                    page_content=content,
                    metadata=metadata
                )
                documents.append(lang_doc)
            
            print(f"✓ DOCX chargé: {os.path.basename(file_path)}")
            
        except Exception as e:
            print(f"✗ Erreur lors du chargement DOCX {file_path}: {e}")
        
        return documents
    
    def load_csv_documents(self, file_path: str) -> List[LangChainDocument]:
        """Charger un fichier CSV"""
        documents = []
        
        try:
            df = pd.read_csv(file_path)
            
            # Convertir chaque ligne en document
            for index, row in df.iterrows():
                content = " | ".join([f"{col}: {str(val)}" for col, val in row.items()])
                
                metadata = {
                    "source": os.path.basename(file_path),
                    "file_path": file_path,
                    "file_type": "csv",
                    "row": index + 1
                }
                
                doc = LangChainDocument(
                    page_content=content,
                    metadata=metadata
                )
                documents.append(doc)
            
            print(f"✓ CSV chargé: {os.path.basename(file_path)} ({len(documents)} lignes)")
            
        except Exception as e:
            print(f"✗ Erreur lors du chargement CSV {file_path}: {e}")
        
        return documents
    
    def load_all_documents(self) -> List[LangChainDocument]:
        """Charger tous les documents du répertoire data/"""
        all_documents = []
        
        if not os.path.exists(self.config.DATA_DIR):
            Logger.error(f"Répertoire de données non trouvé: {self.config.DATA_DIR}")
            return all_documents
        
        Logger.info(f"Recherche de documents dans: {self.config.DATA_DIR}")
        
        for file_path in Path(self.config.DATA_DIR).rglob("*"):
            if file_path.is_file():
                extension = file_path.suffix.lower()
                
                if extension == '.pdf':
                    all_documents.extend(self.load_pdf_documents(str(file_path)))
                elif extension == '.txt':
                    all_documents.extend(self.load_text_documents(str(file_path)))
                elif extension == '.docx':
                    all_documents.extend(self.load_docx_documents(str(file_path)))
                elif extension == '.csv':
                    all_documents.extend(self.load_csv_documents(str(file_path)))
                else:
                    Logger.warning(f"Format non supporté: {file_path}")
        
        Logger.info(f"Total documents chargés: {len(all_documents)}")
        return all_documents
    
    def create_chunks(self, documents: List[LangChainDocument]) -> List[LangChainDocument]:
        """Découper les documents en chunks"""
        Logger.loading("Découpage en chunks...")
        
        chunks = []
        for doc in documents:
            doc_chunks = self.text_splitter.split_documents([doc])
            
            # Pour les PDFs, mettre à jour les métadonnées avec les pages
            if doc.metadata.get("file_type") == "pdf":
                for chunk in doc_chunks:
                    # Essayer de déterminer les pages pour ce chunk
                    file_path = chunk.metadata.get("file_path")
                    if file_path:
                        try:
                            from utils_pdf import create_page_metadata
                            pdf_metadata = create_page_metadata(file_path, chunk.page_content)
                            chunk.metadata.update({
                                "pages": pdf_metadata.get("pages", ["page inconnue"]),
                                "page_range": pdf_metadata.get("page_range", "page inconnue")
                            })
                            print(f"✓ Métadonnées pages ajoutées pour chunk de {os.path.basename(file_path)}")
                        except Exception as e:
                            print(f"⚠️ Erreur extraction pages pour {os.path.basename(file_path)}: {e}")
                            print("   → Utilisation de métadonnées simplifiées")
                            # Continuer avec des métadonnées par défaut
                            chunk.metadata.update({
                                "pages": ["page inconnue"],
                                "page_range": "page inconnue"
                            })
                        except Exception as e:
                            print(f"Erreur métadonnées PDF pour chunk: {e}")
            
            chunks.extend(doc_chunks)
        
        Logger.info(f"Total chunks créés: {len(chunks)}")
        return chunks
    
    def create_vectorstore(self, chunks: List[LangChainDocument]):
        """Créer l'index vectoriel FAISS"""
        if not self.embeddings:
            self._init_embeddings()
        
        Logger.loading("Création de l'index vectoriel FAISS...")
        
        if chunks:
            self.vectorstore = FAISS.from_documents(chunks, self.embeddings)
            
            # Sauvegarder l'index
            os.makedirs(self.config.VECTORSTORE_DIR, exist_ok=True)
            vectorstore_path = os.path.join(self.config.VECTORSTORE_DIR, "faiss_index")
            self.vectorstore.save_local(vectorstore_path)
            
            Logger.success(f"Index vectoriel sauvegardé: {vectorstore_path}")
        else:
            Logger.warning("Aucun chunk à indexer")
    
    def run_ingestion(self):
        """Exécuter le processus complet d'ingestion"""
        print("=== Début de l'ingestion ===")
        
        try:
            # Charger tous les documents
            documents = self.load_all_documents()
            
            if not documents:
                print("Aucun document trouvé. Vérifiez le répertoire data/")
                return False, 0
            
            # Créer les chunks
            chunks = self.create_chunks(documents)
            
            # Créer l'index vectoriel
            self.create_vectorstore(chunks)
            
            # Compter le nombre de fichiers uniques traités
            unique_sources = set()
            for doc in documents:
                if hasattr(doc, 'metadata') and 'source' in doc.metadata:
                    unique_sources.add(doc.metadata['source'])
            
            print("=== Ingestion terminée ===")
            return True, len(unique_sources)
            
        except Exception as e:
            print(f"Erreur lors de l'ingestion: {e}")
            return False, 0

def main():
    """Point d'entrée principal"""
    try:
        Config.ensure_directories()
        
        ingester = DocumentIngester()
        ingester.run_ingestion()
        
    except Exception as e:
        print(f"Erreur lors de l'ingestion: {e}")
        sys.exit(1)

if __name__ == "__main__":
    main()
